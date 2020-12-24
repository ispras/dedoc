import zipfile
from bs4 import BeautifulSoup
import hashlib
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table as DocxTable
from typing import List, Tuple, Optional

from dedoc.extensions import recognized_extensions, recognized_mimes
from dedoc.readers.utils.hierarch_level_extractor import HierarchyLevelExtractor
from dedoc.data_structures.paragraph_metadata import ParagraphMetadata
from dedoc.readers.docx_reader.styles_extractor import StylesExtractor
from dedoc.readers.docx_reader.numbering_extractor import NumberingExtractor
from dedoc.readers.docx_reader.data_structures import Paragraph, ParagraphInfo
from dedoc.structure_parser.heirarchy_level import HierarchyLevel
from dedoc.data_structures.table import Table
from dedoc.data_structures.table_metadata import TableMetadata
from dedoc.data_structures.unstructured_document import UnstructuredDocument
from dedoc.readers.base_reader import BaseReader
from dedoc.data_structures.line_with_meta import LineWithMeta
from dedoc.data_structures.concrete_annotations.bold_annotation import BoldAnnotation
from dedoc.data_structures.concrete_annotations.italic_annotation import ItalicAnnotation
from dedoc.data_structures.concrete_annotations.underlined_annotation import UnderlinedAnnotation
from dedoc.data_structures.concrete_annotations.size_annotation import SizeAnnotation
from dedoc.data_structures.concrete_annotations.alignment_annotation import AlignmentAnnotation
from dedoc.data_structures.concrete_annotations.indentation_annotation import IndentationAnnotation
from dedoc.data_structures.concrete_annotations.style_annotation import StyleAnnotation


class DocxReader(BaseReader):
    def __init__(self):
        self.hierarchy_level_extractor = HierarchyLevelExtractor()
        self.document_xml = None
        self.document_bs_tree = None
        self.paragraph_list = None
        self.path_hash = None
        self.document_bs_tree = None
        self.paragraph_list = None
        self.path_hash = None

    def can_read(self,
                 path: str,
                 mime: str,
                 extension: str,
                 document_type: Optional[str]) -> bool:
        return ((extension in recognized_extensions.docx_like_format or mime in recognized_mimes.docx_like_format) and
                not document_type)

    def read(self,
             path: str,
             document_type: Optional[str] = None,
             parameters: Optional[dict] = None) -> Tuple[UnstructuredDocument, bool]:

        # extract tables
        try:
            document = Document(path)
            tables = [self._process_table(table) for table in document.tables]
        except IndexError:
            tables = []
        except PackageNotFoundError:
            tables = []

        # get hash of document
        with open(path, "rb") as f:
            self.path_hash = hashlib.md5(f.read()).hexdigest()
        # extract text lines
        lines = self._process_lines(path)

        return UnstructuredDocument(lines=lines, tables=tables), True

    @property
    def get_paragraph_list(self) -> List[BeautifulSoup]:
        return self.paragraph_list

    @property
    def get_document_bs_tree(self) -> BeautifulSoup:
        return self.document_bs_tree

    @staticmethod
    def _process_table(table: DocxTable) -> Table:
        cells = [[cell.text for cell in row.cells] for row in table.rows]
        metadata = TableMetadata(page_id=None)
        return Table(cells=cells, metadata=metadata)

    def _process_lines(self,
                       path: str) -> List[LineWithMeta]:
        """
        :param path: path to file for parsing
        :return: list of document lines with annotations
        """
        self.document_xml = zipfile.ZipFile(path)
        self.document_bs_tree = self.__get_bs_tree('word/document.xml')
        self.paragraph_list = []
        if self.document_bs_tree:
            body = self.document_bs_tree.body
        else:
            return []

        styles_extractor = StylesExtractor(self.__get_bs_tree('word/styles.xml'))
        num_tree = self.__get_bs_tree('word/numbering.xml')
        if num_tree:
            numbering_extractor = NumberingExtractor(num_tree, styles_extractor)
        else:
            numbering_extractor = None
        styles_extractor.numbering_extractor = numbering_extractor

        footers, headers = [], []
        for i in range(1, 4):
            footer = self.__get_bs_tree('word/footer' + str(i) + '.xml')
            if footer:
                footers.append(footer)
            header = self.__get_bs_tree('word/header' + str(i) + '.xml')
            if header:
                headers.append(header)
        footnotes = self.__get_bs_tree('word/footnotes.xml')
        endnotes = self.__get_bs_tree('word/endnotes.xml')

        # the list of paragraph with their properties
        paragraph_list = []
        for header in headers:
            self.__add_to_paragraph_list(header)

        for paragraph in body:
            # ignore tables
            if paragraph.name == 'tbl':
                continue
            if paragraph.name != 'p':
                self.__add_to_paragraph_list(paragraph)
                continue
            self.paragraph_list.append(paragraph)

        if footnotes:
            self.__add_to_paragraph_list(footnotes)
        if endnotes:
            self.__add_to_paragraph_list(endnotes)
        for footer in footers:
            self.__add_to_paragraph_list(footer)

        paragraph_list = []
        for paragraph in self.paragraph_list:
            paragraph_list.append(Paragraph(paragraph, styles_extractor, numbering_extractor))

        return self._get_lines_with_meta(paragraph_list)

    def __get_bs_tree(self,
                      filename: str) -> Optional[BeautifulSoup]:
        """
        gets xml bs tree from the given file inside the self.document_xml
        :param filename: name of file to extract the tree
        :return: BeautifulSoup tree or None if file wasn't found
        """
        try:
            tree = BeautifulSoup(self.document_xml.read(filename), 'xml')
        except KeyError:
            tree = None
        return tree

    def __add_to_paragraph_list(self, tree: BeautifulSoup) -> None:
        self.paragraph_list += tree.find_all('w:p')

    def _get_lines_with_meta(self,
                             paragraph_list: List[Paragraph]) -> List[LineWithMeta]:
        """
        :param paragraph_list: list of Paragraph
        :return: list of LineWithMeta
        """
        lines_with_meta = []
        paragraph_id = 0

        for paragraph in paragraph_list:

            # line with meta:
            # {"text": "",
            #  "type": ""("paragraph" ,"list_item", "raw_text", "style_header"),
            #  "level": (1,1) or None (hierarchy_level),
            #  "annotations": [["size", start, end, size], ["bold", start, end, "True"], ...]}
            paragraph_properties = ParagraphInfo(paragraph)
            line_with_meta = paragraph_properties.get_info()

            text = line_with_meta["text"]
            uid = '{}_{}'.format(self.path_hash, line_with_meta["uid"])

            paragraph_type = line_with_meta["type"]
            level = line_with_meta["level"]
            if level:
                hierarchy_level = HierarchyLevel(level[0], level[1], False, paragraph_type)
            else:
                hierarchy_level = HierarchyLevel(None, None, False, "raw_text")

            dict2annotations = {
                "bold": BoldAnnotation,
                "italic": ItalicAnnotation,
                "underlined": UnderlinedAnnotation,
                "size": SizeAnnotation,
                "indentation": IndentationAnnotation,
                "alignment": AlignmentAnnotation,
                "style": StyleAnnotation
            }
            annotations = []
            for annotation in line_with_meta["annotations"]:
                annotations.append(dict2annotations[annotation[0]](*annotation[1:]))

            paragraph_id += 1
            metadata = ParagraphMetadata(paragraph_type=paragraph_type,
                                         predicted_classes=None,
                                         page_id=0,
                                         line_id=paragraph_id)

            lines_with_meta.append(LineWithMeta(line=text,
                                                hierarchy_level=hierarchy_level,
                                                metadata=metadata,
                                                annotations=annotations,
                                                uid=uid))
            lines_with_meta = self.hierarchy_level_extractor.get_hierarchy_level(lines_with_meta)
        return lines_with_meta
